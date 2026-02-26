"""
AI Resume & Portfolio Builder
Flask web app that collects user data, calls HuggingFace Inference API (Llama-3.1-8B-Instruct),
and generates Resume (DOCX), Cover Letter (DOCX), and Portfolio (HTML)
bundled as a ZIP download.
"""

import io
import json
import os
import re
import zipfile
import requests
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv

load_dotenv()  # loads HF_TOKEN from .env file

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# HuggingFace Inference API — Llama-3.1-8B-Instruct
# Free tier — get your key at: https://huggingface.co/settings/tokens
# ─────────────────────────────────────────────────────────────────────────────
HF_ROUTER_URL = "https://router.huggingface.co/v1/chat/completions"
HF_MODEL      = "meta-llama/Llama-3.1-8B-Instruct"   # free via hf-inference provider
HF_TOKEN      = os.environ.get("HF_TOKEN", "")  # Set this in your .env file


# ─────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    try:
        data = request.get_json()
        # Validate token is configured
        if not HF_TOKEN:
            return jsonify({"error": "HF_TOKEN not set. Please add it to your .env file."}), 500

        # 1. Call HuggingFace Inference API
        ai_content = call_huggingface(data)

        # 2. Build documents in memory
        resume_bytes   = build_resume_docx(data, ai_content)
        cover_bytes    = build_cover_letter_docx(data, ai_content)
        portfolio_html = build_portfolio_html(data, ai_content)

        # 3. Bundle into ZIP
        name_slug = data.get("name", "Resume").replace(" ", "_")
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{name_slug}_Resume.docx",       resume_bytes)
            zf.writestr(f"{name_slug}_Cover_Letter.docx", cover_bytes)
            zf.writestr(f"{name_slug}_Portfolio.html",    portfolio_html)
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"{name_slug}_Career_Documents.zip"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─────────────────────────────────────────────
# HUGGINGFACE INFERENCE API
# ─────────────────────────────────────────────

def build_prompt(data: dict) -> str:
    """Build the user prompt for the chat completions API."""
    is_fresher = data.get("experience_level") == "fresher"
    has_jd = bool(data.get("job_description", "").strip())

    exp_section = ""
    if not is_fresher:
        exp_lines = [
            f"  - {e['role']} at {e['company']} ({e['duration']}): {e['description']}"
            for e in data.get("experiences", []) if e.get("company")
        ]
        if exp_lines:
            exp_section = "WORK EXPERIENCE:\n" + "\n".join(exp_lines)

    projects_text = "\n".join([
        f"  {i+1}. {p['name']}: {p['description']} | Tech: {p['tech']} | Demo: {p.get('demo','')} | GitHub: {p.get('github','')}"
        for i, p in enumerate(data.get("projects", [])) if p.get("name")
    ]) or "  None provided"

    certs_text = "\n".join([
        f"  - {c['name']} by {c.get('issuer','')} ({c.get('year','')})"
        for c in data.get("certifications", []) if c.get("name")
    ]) or "  None"

    jd_instruction = (
        f"JOB DESCRIPTION (tailor resume keywords to match this):\n{data['job_description']}"
        if has_jd else
        "No JD provided — create a strong general professional resume."
    )

    parts = [
        "Generate ATS-optimized resume content for this candidate.",
        "Return ONLY a valid JSON object with no markdown, no code fences, no extra text.",
        "",
        f"Name: {data.get('name')} | Email: {data.get('email')} | Phone: {data.get('phone')}",
        f"Location: {data.get('location')} | LinkedIn: {data.get('linkedin','')} | GitHub: {data.get('github','')}",
        f"Primary Degree: {data.get('degree')} in {data.get('field')} - {data.get('university')}, {data.get('uni_location')} ({data.get('graduation_year')}) CGPA: {data.get('cgpa')}",
        f"Second Degree: {data.get('degree2','')} in {data.get('field2','')} - {data.get('university2','')} ({data.get('graduation_year2','')}) CGPA: {data.get('cgpa2','')}".strip('" -') if data.get('degree2') else "",
        f"Technical Skills: {data.get('technical_skills')}",
        f"Soft Skills: {data.get('soft_skills','')}",
        "",
        "Projects:", projects_text, "",
        "Certifications:", certs_text, "",
        f"Experience Level: {data.get('experience_level')}",
        exp_section, "",
        f"Target Role: {data.get('job_role', 'Software Engineer')}",
        jd_instruction, "",
        "Rules:",
        "- Professional Summary: Write 4 to 5 detailed sentences. Mention degree, years of experience or fresher status, top 3-4 technical skills, key strengths, and career goal for the target role.",
        "- Use strong action verbs (Built/Developed/Led/Optimized). Quantify results wherever possible.",
        "- Group skills by category (Languages, Frameworks, Databases, Tools/Cloud).",
        "- Cover letter: exactly 3 formal paragraphs, compelling and specific.",
        "",
        'Return this exact JSON (nothing else):',
        '{',
        '  "summary": "4-5 detailed sentences: mention degree, skills, experience level, strengths, and career goal",',
        '  "skill_bullets": ["Languages: Python, JS", "Frameworks: React, Flask"],',
        '  "experience_bullets": [{"role": "...", "company": "...", "duration": "...", "bullets": ["achievement"]}],',
        '  "project_bullets": [{"name": "...", "bullets": ["impact 1", "impact 2"]}],',
        '  "education_notes": ["CGPA: X.XX", "Coursework: DSA, OS"],',
        '  "cover_letter_subject": "Application for Role Position",',
        '  "cover_letter_body": "Para 1.\\n\\nPara 2.\\n\\nPara 3.",',
        '  "portfolio_tagline": "Short tagline here",',
        '  "portfolio_about": "About para 1.\\n\\nAbout para 2.",',
        '  "ats_keywords": ["keyword1", "keyword2"]',
        '}',
    ]
    return "\n".join(parts)


def extract_json(text: str) -> dict:
    """Robustly extract JSON from the model's raw output."""
    text = text.strip()
    # Remove markdown fences
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    # Direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # Find first { ... } block
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    # Fix trailing commas and retry
    cleaned = re.sub(r",\s*([}\]])", r"\1", text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    raise ValueError(
        f"Could not parse JSON from model response.\n"
        f"First 500 chars of output:\n{text[:500]}"
    )


def call_huggingface(data: dict) -> dict:
    """
    Call HuggingFace router (OpenAI-compatible chat completions API).
    Endpoint: https://router.huggingface.co/v1/chat/completions
    Model: meta-llama/Llama-3.1-8B-Instruct (free via hf-inference provider)
    Token: Set HF_TOKEN constant at the top of this file.
    """
    system_msg = (
        "You are an expert resume writer. Always respond with ONLY valid JSON. "
        "No markdown, no code fences, no explanation. Pure JSON object only. "
        "For the summary field, always write 4 to 5 detailed sentences covering degree, skills, experience, and career goals."
    )
    user_msg = build_prompt(data)

    headers = {
        "Authorization": f"Bearer {HF_TOKEN}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": HF_MODEL,
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user",   "content": user_msg},
        ],
        "max_tokens": 2000,
        "temperature": 0.3,
        "top_p": 0.9,
        "stream": False,
    }

    resp = requests.post(HF_ROUTER_URL, headers=headers, json=payload, timeout=120)

    if resp.status_code == 401:
        raise ValueError(
            "Invalid HuggingFace token. "
            "Get your free key at: huggingface.co/settings/tokens"
        )
    elif resp.status_code == 403:
        raise ValueError(
            "Access denied. Make sure your token has 'Make calls to Inference Providers' "
            "permission enabled at: huggingface.co/settings/tokens"
        )
    elif resp.status_code == 404:
        raise ValueError(
            "Model not found on HuggingFace router. "
            f"Check model name: {HF_MODEL}"
        )
    elif resp.status_code == 429:
        raise ValueError(
            "Rate limit exceeded. Please wait a moment and try again."
        )
    elif resp.status_code != 200:
        raise ValueError(
            f"HuggingFace API error {resp.status_code}: {resp.text[:400]}"
        )

    result = resp.json()

    # OpenAI-compatible response format
    try:
        generated_text = result["choices"][0]["message"]["content"]
    except (KeyError, IndexError):
        raise ValueError(f"Unexpected response format: {str(result)[:400]}")

    if not generated_text or not generated_text.strip():
        raise ValueError("HuggingFace returned an empty response. Please try again.")

    return extract_json(generated_text)


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────

def add_section_heading(doc, title, color_hex="1F4E79"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    return p


def add_horizontal_rule(doc, color="1F4E79"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p


# ─────────────────────────────────────────────
# RESUME DOCX
# ─────────────────────────────────────────────

def build_resume_docx(data: dict, ai: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Name
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_after = Pt(2)
    nr = np_.add_run(data.get("name", ""))
    nr.bold = True; nr.font.size = Pt(22); nr.font.name = "Calibri"
    nr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Job role
    if data.get("job_role"):
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.paragraph_format.space_after = Pt(3)
        rr = rp.add_run(data["job_role"])
        rr.font.size = Pt(11); rr.font.name = "Calibri"
        rr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # Contact
    parts = [v for v in [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github"), data.get("website")
    ] if v]
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    cr = cp.add_run("  |  ".join(parts))
    cr.font.size = Pt(9); cr.font.name = "Calibri"
    cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Summary
    if ai.get("summary"):
        add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ai["summary"])
        r.font.size = Pt(10); r.font.name = "Calibri"

    # Education
    add_section_heading(doc, "Education")

    def add_degree_block(deg, field, uni, uni_loc, grad_yr, cgpa_val, notes=None):
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(4); ep.paragraph_format.space_after = Pt(0)
        r1 = ep.add_run(f"{deg} in {field}")
        r1.bold = True; r1.font.size = Pt(10.5); r1.font.name = "Calibri"
        r2 = ep.add_run(f"   {grad_yr}")
        r2.font.size = Pt(10); r2.font.name = "Calibri"
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        up = doc.add_paragraph()
        up.paragraph_format.space_before = Pt(0); up.paragraph_format.space_after = Pt(2)
        ur = up.add_run(f"{uni}, {uni_loc}")
        ur.font.size = Pt(10); ur.font.name = "Calibri"; ur.italic = True
        ur.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        if cgpa_val:
            add_bullet(doc, f"CGPA: {cgpa_val}")
        if notes:
            for note in notes:
                if note.strip():
                    add_bullet(doc, note)

    # Primary degree
    add_degree_block(
        data.get('degree',''), data.get('field',''),
        data.get('university',''), data.get('uni_location',''),
        data.get('graduation_year',''), data.get('cgpa',''),
        ai.get("education_notes", [])
    )
    # Second degree (if provided)
    if data.get('degree2'):
        add_degree_block(
            data.get('degree2',''), data.get('field2',''),
            data.get('university2',''), data.get('uni_location2',''),
            data.get('graduation_year2',''), data.get('cgpa2','')
        )

    # Skills
    add_section_heading(doc, "Technical Skills")
    for bullet in ai.get("skill_bullets", [data.get("technical_skills", "")]):
        if bullet.strip():
            add_bullet(doc, bullet)

    # Experience
    is_fresher = data.get("experience_level") == "fresher"
    exp_bullets = ai.get("experience_bullets", [])
    if not is_fresher and exp_bullets:
        add_section_heading(doc, "Work Experience")
        for exp in exp_bullets:
            if not exp.get("company"):
                continue
            ep2 = doc.add_paragraph()
            ep2.paragraph_format.space_before = Pt(4); ep2.paragraph_format.space_after = Pt(0)
            rr1 = ep2.add_run(exp.get("role", ""))
            rr1.bold = True; rr1.font.size = Pt(10.5); rr1.font.name = "Calibri"
            rr2 = ep2.add_run(f"   {exp.get('duration', '')}")
            rr2.font.size = Pt(9.5); rr2.font.name = "Calibri"
            rr2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            cp2 = doc.add_paragraph()
            cp2.paragraph_format.space_before = Pt(0); cp2.paragraph_format.space_after = Pt(1)
            cr2 = cp2.add_run(exp.get("company", ""))
            cr2.font.size = Pt(10); cr2.font.name = "Calibri"; cr2.italic = True
            cr2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            for b in exp.get("bullets", []):
                if b.strip():
                    add_bullet(doc, b)

    # Projects
    proj_bullets = ai.get("project_bullets", [])
    if proj_bullets:
        add_section_heading(doc, "Projects")
        for proj in proj_bullets:
            if not proj.get("name"):
                continue
            pp = doc.add_paragraph()
            pp.paragraph_format.space_before = Pt(4); pp.paragraph_format.space_after = Pt(0)
            matching = next((p for p in data.get("projects", [])
                             if p.get("name","").strip() == proj["name"].strip()), {})
            link_parts = []
            if matching.get("demo"):   link_parts.append(f"Demo: {matching['demo']}")
            if matching.get("github"): link_parts.append(f"GitHub: {matching['github']}")
            pr = pp.add_run(proj["name"])
            pr.bold = True; pr.font.size = Pt(10.5); pr.font.name = "Calibri"
            if link_parts:
                pr2 = pp.add_run(f"   |  {' | '.join(link_parts)}")
                pr2.font.size = Pt(9); pr2.font.name = "Calibri"
                pr2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            for b in proj.get("bullets", []):
                if b.strip():
                    add_bullet(doc, b)

    # Certifications
    certs = [c for c in data.get("certifications", []) if c.get("name")]
    if certs:
        add_section_heading(doc, "Certifications")
        for c in certs:
            parts = [c["name"]]
            if c.get("issuer"): parts.append(c["issuer"])
            if c.get("year"):   parts.append(c["year"])
            add_bullet(doc, " — ".join(parts))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# COVER LETTER DOCX
# ─────────────────────────────────────────────

def build_cover_letter_docx(data: dict, ai: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def normal(text, bold=False, size=11, color=None, space_after=8, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = "Calibri"
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    normal(data.get("name", ""), bold=True, size=13, color=(0x1F, 0x4E, 0x79))
    contact = "  |  ".join(filter(None, [data.get("email"), data.get("phone"), data.get("location")]))
    normal(contact, size=10, color=(0x55, 0x55, 0x55))
    if data.get("linkedin"):
        normal(data["linkedin"], size=10, color=(0x44, 0x72, 0xC4))

    add_horizontal_rule(doc)

    normal(datetime.today().strftime("%B %d, %Y"), size=11, space_after=12)
    normal("Hiring Manager / Recruitment Team", size=11)
    subject = ai.get("cover_letter_subject", f"Application for {data.get('job_role','the Position')}")
    normal(f"Re: {subject}", bold=True, size=11, space_after=14)
    normal("Dear Hiring Manager,", size=11, space_after=10)

    body = ai.get("cover_letter_body", "")
    for para in body.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(para)
            r.font.size = Pt(11); r.font.name = "Calibri"

    normal("Sincerely,", size=11, space_after=28)
    normal(data.get("name", ""), bold=True, size=11, color=(0x1F, 0x4E, 0x79))
    normal("  |  ".join(filter(None, [data.get("email"), data.get("phone")])),
           size=10, color=(0x55, 0x55, 0x55))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# PORTFOLIO HTML
# ─────────────────────────────────────────────

def build_portfolio_html(data: dict, ai: dict) -> str:
    name        = data.get("name", "")
    job_role    = data.get("job_role", "Developer")
    email       = data.get("email", "")
    phone       = data.get("phone", "")
    location    = data.get("location", "")
    linkedin    = data.get("linkedin", "")
    github      = data.get("github", "")
    website     = data.get("website", "")
    tech_skills = [s.strip() for s in data.get("technical_skills", "").split(",") if s.strip()]
    projects    = [p for p in data.get("projects", []) if p.get("name")]
    certs       = [c for c in data.get("certifications", []) if c.get("name")]
    experiences = data.get("experiences", [])
    is_fresher  = data.get("experience_level") == "fresher"
    tagline     = ai.get("portfolio_tagline", "Building the Future, One Line at a Time")
    about_raw   = ai.get("portfolio_about", "")
    about_html  = "".join(f"<p>{p.strip()}</p>" for p in about_raw.split("\n") if p.strip())
    first_name  = name.split()[0] if name else "Dev"

    skills_html = "\n".join(f'<span class="skill-chip">{s}</span>' for s in tech_skills)

    proj_cards = ""
    for p in projects:
        tech_tags = "".join(
            f'<span class="tech-tag">{t.strip()}</span>'
            for t in p.get("tech","").split(",") if t.strip()
        )
        links = ""
        if p.get("demo"):   links += f'<a href="{p["demo"]}" class="proj-link" target="_blank">🔗 Live Demo</a>'
        if p.get("github"): links += f'<a href="https://{p["github"]}" class="proj-link" target="_blank">💻 GitHub</a>'
        proj_cards += f"""
        <div class="project-card">
          <div class="proj-title">{p['name']}</div>
          <p class="proj-desc">{p.get('description','')}</p>
          <div class="tech-tags">{tech_tags}</div>
          <div class="proj-links">{links}</div>
        </div>"""

    exp_html = ""
    if not is_fresher:
        for e in experiences:
            if e.get("company"):
                exp_html += f"""
        <div class="exp-card">
          <div class="exp-top"><span class="exp-role">{e.get('role','')}</span><span class="exp-dur">{e.get('duration','')}</span></div>
          <div class="exp-company">{e.get('company','')}</div>
          <p class="exp-desc">{e.get('description','')}</p>
        </div>"""

    cert_html = ""
    for c in certs:
        cert_html += f"""
        <div class="cert-card">
          <div class="cert-name">{c['name']}</div>
          <div class="cert-issuer">{c.get('issuer','')}</div>
          <div class="cert-year">{c.get('year','')}</div>
        </div>"""

    contact_links = ""
    if email:    contact_links += f'<a href="mailto:{email}" class="btn btn-filled">✉️ Email Me</a>'
    if linkedin: contact_links += f'<a href="https://{linkedin}" class="btn btn-outline" target="_blank">LinkedIn</a>'
    if github:   contact_links += f'<a href="https://{github}" class="btn btn-outline" target="_blank">GitHub</a>'
    if phone:    contact_links += f'<a href="tel:{phone}" class="btn btn-outline">📞 Call</a>'

    hero_links = ""
    if github:   hero_links += f'<a href="https://{github}" class="btn btn-filled" target="_blank">View GitHub →</a>'
    if linkedin: hero_links += f'<a href="https://{linkedin}" class="btn btn-outline" target="_blank">LinkedIn</a>'
    if email:    hero_links += f'<a href="mailto:{email}" class="btn btn-outline">Contact Me</a>'

    nav_exp  = '<li><a href="#experience">Experience</a></li>' if not is_fresher and experiences else ""
    nav_cert = '<li><a href="#certifications">Certs</a></li>' if certs else ""

    sec_exp = f"""
    <section id="experience">
      <p class="sec-label">Where I've Worked</p>
      <h2>Experience</h2>
      <div class="exp-list">{exp_html}</div>
    </section>""" if not is_fresher and exp_html else ""

    sec_cert = f"""
    <section id="certifications">
      <p class="sec-label">Credentials</p>
      <h2>Certifications</h2>
      <div class="cert-grid">{cert_html}</div>
    </section>""" if cert_html else ""

    about_info = ""
    for label, value, href in [
        ("Email",    email,    f"mailto:{email}"),
        ("Phone",    phone,    f"tel:{phone}"),
        ("Location", location, ""),
        ("LinkedIn", linkedin, f"https://{linkedin}"),
        ("GitHub",   github,   f"https://{github}"),
        ("Website",  website,  f"https://{website}"),
    ]:
        if value:
            display = f'<a href="{href}" style="color:var(--accent)">{value}</a>' if href else value
            about_info += f'<div class="info-item"><span class="info-label">{label}</span><span class="info-val">{display}</span></div>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/><meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>{name} | Portfolio</title>
  <link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;600;700&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&display=swap" rel="stylesheet"/>
  <style>
    *,*::before,*::after{{box-sizing:border-box;margin:0;padding:0;}}
    :root{{--bg:#020917;--surface:#0d1424;--border:#1e2d45;--text:#e8f0fe;--sub:#7b92b4;--accent:#4f8ef7;--accent2:#2563eb;--green:#22c55e;}}
    html{{scroll-behavior:smooth;}}
    body{{background:var(--bg);color:var(--text);font-family:'DM Sans',sans-serif;line-height:1.6;}}
    ::-webkit-scrollbar{{width:5px;}};::-webkit-scrollbar-track{{background:var(--bg);}};::-webkit-scrollbar-thumb{{background:var(--border);border-radius:4px;}}
    nav{{position:fixed;top:0;width:100%;background:rgba(2,9,23,0.85);backdrop-filter:blur(14px);border-bottom:1px solid var(--border);padding:14px 6%;display:flex;justify-content:space-between;align-items:center;z-index:999;}}
    .nav-logo{{font-family:'Space Grotesk';font-weight:700;font-size:18px;color:var(--text);text-decoration:none;}}
    .nav-logo span{{color:var(--accent);}}
    nav ul{{list-style:none;display:flex;gap:28px;}}
    nav a{{color:var(--sub);text-decoration:none;font-size:14px;font-weight:500;transition:color 0.2s;}}
    nav a:hover{{color:var(--text);}}
    section{{padding:90px 8%;}}
    .sec-label{{font-size:11px;text-transform:uppercase;letter-spacing:3px;color:var(--accent);font-weight:700;margin-bottom:10px;}}
    h2{{font-family:'Space Grotesk';font-size:clamp(26px,3vw,36px);font-weight:700;margin-bottom:40px;}}
    .hero{{min-height:100vh;padding-top:120px;display:flex;align-items:center;background:radial-gradient(ellipse 60% 50% at 10% 60%,rgba(79,142,247,0.10) 0%,transparent 70%),radial-gradient(ellipse 40% 40% at 90% 30%,rgba(37,99,235,0.08) 0%,transparent 60%),var(--bg);}}
    .hero-inner{{max-width:720px;}}
    .badge{{display:inline-flex;align-items:center;gap:8px;background:rgba(34,197,94,0.12);border:1px solid rgba(34,197,94,0.3);color:var(--green);padding:5px 14px;border-radius:20px;font-size:13px;font-weight:500;margin-bottom:24px;}}
    .badge::before{{content:'';width:8px;height:8px;border-radius:50%;background:var(--green);animation:pulse 2s infinite;}}
    @keyframes pulse{{0%,100%{{opacity:1;}}50%{{opacity:0.4;}}}}
    h1{{font-family:'Space Grotesk';font-size:clamp(38px,6vw,68px);font-weight:700;line-height:1.08;margin-bottom:16px;}}
    h1 .hl{{color:var(--accent);}}
    .tagline{{font-size:clamp(16px,2vw,20px);color:var(--sub);margin-bottom:20px;font-weight:300;}}
    .hero-btns{{display:flex;gap:14px;flex-wrap:wrap;margin-top:28px;}}
    .btn{{padding:12px 28px;border-radius:8px;font-weight:600;text-decoration:none;font-size:14px;transition:all 0.2s;display:inline-block;border:none;cursor:pointer;font-family:inherit;}}
    .btn-filled{{background:linear-gradient(135deg,var(--accent),var(--accent2));color:white;}}
    .btn-filled:hover{{transform:translateY(-2px);box-shadow:0 8px 28px rgba(79,142,247,0.35);}}
    .btn-outline{{border:1px solid var(--border);color:var(--text);background:transparent;}}
    .btn-outline:hover{{border-color:var(--accent);color:var(--accent);}}
    #about{{background:var(--surface);border-top:1px solid var(--border);border-bottom:1px solid var(--border);}}
    .about-grid{{display:grid;grid-template-columns:1fr 1fr;gap:48px;}}
    .about-text p{{color:var(--sub);font-size:15px;line-height:1.8;margin-bottom:14px;}}
    .info-item{{display:flex;gap:10px;margin-bottom:12px;font-size:14px;}}
    .info-label{{color:var(--accent);font-weight:600;min-width:80px;}}
    .info-val{{color:var(--sub);}}
    .skills-wrap{{display:flex;flex-wrap:wrap;gap:10px;}}
    .skill-chip{{background:var(--surface);border:1px solid var(--border);border-radius:8px;padding:8px 18px;font-size:13px;font-weight:500;transition:all 0.2s;}}
    .skill-chip:hover{{border-color:var(--accent);color:var(--accent);transform:translateY(-2px);}}
    .projects-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(300px,1fr));gap:20px;}}
    .project-card{{background:var(--surface);border:1px solid var(--border);border-radius:14px;padding:24px;transition:all 0.25s;}}
    .project-card:hover{{border-color:var(--accent);transform:translateY(-4px);box-shadow:0 12px 40px rgba(79,142,247,0.12);}}
    .proj-title{{font-family:'Space Grotesk';font-size:17px;font-weight:700;margin-bottom:8px;}}
    .proj-desc{{color:var(--sub);font-size:13px;line-height:1.6;margin-bottom:14px;}}
    .tech-tags{{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:14px;}}
    .tech-tag{{background:rgba(79,142,247,0.12);color:#79b8ff;padding:2px 10px;border-radius:12px;font-size:12px;}}
    .proj-links{{display:flex;gap:14px;}}
    .proj-link{{color:var(--accent);font-size:13px;text-decoration:none;font-weight:500;}}
    .proj-link:hover{{text-decoration:underline;}}
    .edu-card{{background:var(--surface);border:1px solid var(--border);border-radius:14px;padding:28px;max-width:580px;}}
    .edu-deg{{font-family:'Space Grotesk';font-size:19px;font-weight:700;margin-bottom:6px;}}
    .edu-uni{{color:var(--accent);font-size:14px;margin-bottom:8px;}}
    .edu-meta{{color:var(--sub);font-size:13px;}}
    .exp-list{{display:flex;flex-direction:column;gap:16px;}}
    .exp-card{{background:var(--surface);border:1px solid var(--border);border-radius:12px;padding:22px;}}
    .exp-top{{display:flex;justify-content:space-between;align-items:baseline;margin-bottom:4px;}}
    .exp-role{{font-weight:700;font-size:16px;}}
    .exp-dur{{color:var(--sub);font-size:13px;}}
    .exp-company{{color:var(--accent);font-size:13px;margin-bottom:8px;}}
    .exp-desc{{color:var(--sub);font-size:14px;line-height:1.6;}}
    .cert-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:14px;}}
    .cert-card{{background:var(--surface);border:1px solid var(--border);border-radius:10px;padding:18px;}}
    .cert-name{{font-weight:700;font-size:14px;margin-bottom:4px;}}
    .cert-issuer{{color:var(--accent);font-size:12px;}}
    .cert-year{{color:var(--sub);font-size:12px;margin-top:4px;}}
    .contact-box{{background:linear-gradient(135deg,var(--surface),rgba(79,142,247,0.05));border:1px solid var(--border);border-radius:18px;padding:56px;text-align:center;max-width:680px;margin:0 auto;}}
    .contact-sub{{color:var(--sub);font-size:16px;margin-bottom:30px;}}
    .contact-btns{{display:flex;justify-content:center;gap:14px;flex-wrap:wrap;}}
    footer{{border-top:1px solid var(--border);padding:24px 8%;text-align:center;color:var(--sub);font-size:13px;}}
    @media(max-width:700px){{.about-grid{{grid-template-columns:1fr;}}section{{padding:70px 5%;}}nav ul{{gap:16px;}}}}
  </style>
</head>
<body>
<nav>
  <a href="#home" class="nav-logo">{first_name}<span>.</span></a>
  <ul>
    <li><a href="#about">About</a></li>
    <li><a href="#skills">Skills</a></li>
    <li><a href="#projects">Projects</a></li>
    {nav_exp}
    <li><a href="#education">Education</a></li>
    {nav_cert}
    <li><a href="#contact">Contact</a></li>
  </ul>
</nav>

<section class="hero" id="home">
  <div class="hero-inner">
    <span class="badge">Available for Opportunities</span>
    <h1>Hi, I'm <span class="hl">{name}</span></h1>
    <p class="tagline">{tagline}</p>
    <p style="color:var(--sub);font-size:16px;">{job_role} &mdash; {location}</p>
    <div class="hero-btns">{hero_links}</div>
  </div>
</section>

<section id="about">
  <p class="sec-label">Who I Am</p><h2>About Me</h2>
  <div class="about-grid">
    <div class="about-text">{about_html}</div>
    <div class="about-info">{about_info}</div>
  </div>
</section>

<section id="skills">
  <p class="sec-label">What I Know</p><h2>Skills &amp; Technologies</h2>
  <div class="skills-wrap">{skills_html}</div>
</section>

<section id="projects">
  <p class="sec-label">What I've Built</p><h2>Projects</h2>
  <div class="projects-grid">{proj_cards}</div>
</section>

{sec_exp}

<section id="education">
  <p class="sec-label">Academic Background</p><h2>Education</h2>
  <div class="edu-card">
    <p class="edu-deg">{data.get('degree')} in {data.get('field')}</p>
    <p class="edu-uni">{data.get('university')}, {data.get('uni_location')}</p>
    <p class="edu-meta">Graduated: {data.get('graduation_year')} &nbsp;|&nbsp; CGPA: {data.get('cgpa')}</p>
  </div>
  {"<div class='edu-card' style='margin-top:14px'><p class='edu-deg'>" + data.get('degree2') + " in " + data.get('field2','') + "</p><p class='edu-uni'>" + data.get('university2','') + ", " + data.get('uni_location2','') + "</p><p class='edu-meta'>Graduated: " + data.get('graduation_year2','') + " &nbsp;|&nbsp; CGPA: " + data.get('cgpa2','') + "</p></div>" if data.get('degree2') else ""}
</section>

{sec_cert}

<section id="contact">
  <div class="contact-box">
    <p class="sec-label" style="text-align:center">Let's Connect</p>
    <h2>Get In Touch</h2>
    <p class="contact-sub">Open to opportunities, collaborations, and interesting conversations.</p>
    <div class="contact-btns">{contact_links}</div>
  </div>
</section>

<footer><p>Crafted by <strong>{name}</strong> &mdash; {location}</p></footer>
</body>
</html>"""


# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  AI Resume & Portfolio Builder  (HuggingFace Edition)")
    print("  Model : Llama-3.1-8B-Instruct  (Free via HF Router)")
    print("  Open  : http://localhost:5000")
    print("=" * 60)
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)
